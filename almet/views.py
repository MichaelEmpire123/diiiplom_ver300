import io
import os
from django.views.decorators.csrf import csrf_exempt
import re
import secrets
import string
from datetime import datetime
import pandas as pd
from asgiref.sync import async_to_sync
from django.contrib import messages
from django.contrib.auth.hashers import make_password
from django.core.exceptions import PermissionDenied
from django.http import HttpResponse, JsonResponse, HttpResponseForbidden, Http404
from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth import authenticate, login, logout, update_session_auth_hash
from django.contrib.auth.decorators import login_required, user_passes_test
from django.template.loader import render_to_string
from django.utils import timezone
from django.utils.timezone import make_aware
from docx import Document
from docx.shared import Inches
from .forms import AppealForm, MessageForm, ServiceForm, EmployeeRegistrationForm, Edit_AppealForm, ReportForm
from .models import User, Citizen, Street, City, Status, Appeals, Message, Processing_appeals, Category, Sotrudniki, \
    Service
from django.db.models import Max, OuterRef, Subquery, Q, Count
import xml.etree.ElementTree as ET
from django.utils.timezone import now
import json
from channels.layers import get_channel_layer

def validate_cyrillic(text):
    """Проверяет, содержит ли текст только кириллические буквы, пробелы и дефисы"""
    return bool(re.fullmatch(r'^[А-Яа-яЁё\s-]+$', text))

def index(request):
    # Получаем все службы
    services = Service.objects.all()
    return render(request, 'almet/index.html', {'services': services})


def login_view(request):
    if request.user.is_authenticated:
        return redirect('profile')  # Замените 'profile' на нужный URL


    if request.method == 'GET':
        # Рендерим страницу авторизации для GET-запросов
        return render(request, 'auth/login.html')


    if request.method == 'POST':
        # Получаем данные из формы
        login_input = request.POST.get('login_input')  # Поле для email или телефона
        password = request.POST.get('password')

        # Проверяем, что все поля заполнены
        if not login_input or not password:
            return JsonResponse({'success': False, 'message': 'Пожалуйста, заполните все поля.'})

        user = None

        # Проверяем, является ли введённое значение email
        if '@' in login_input:
            # Если введён email, проверяем как жителя, так и сотрудника
            try:
                # Ищем пользователя по email
                user = User.objects.get(email=login_input)

                # Проверяем пароль вручную
                if user.check_password(password):
                    login(request, user)
                    return JsonResponse({'success': True, 'message': 'Авторизация успешна!', 'redirect': '/profile/'})
                else:
                    return JsonResponse({'success': False, 'message': 'Неверный email/телефон или пароль.'})
            except User.DoesNotExist:
                return JsonResponse({'success': False, 'message': 'Неверный email/телефон или пароль.'})
        else:
            # Если введён телефон, проверяем только жителя
            try:
                citizen = Citizen.objects.get(tel=login_input)
                user = authenticate(request, email=citizen.email, password=password)

                if user is not None:
                    login(request, user)
                    return JsonResponse({'success': True, 'message': 'Авторизация успешна!', 'redirect': '/profile/'})
                else:
                    return JsonResponse({'success': False, 'message': 'Неверный email/телефон или пароль.'})
            except Citizen.DoesNotExist:
                return JsonResponse({'success': False, 'message': 'Неверный email/телефон или пароль.'})

    return JsonResponse({'success': False, 'message': 'Недопустимый метод запроса.'})




def register_view(request):
    if request.user.is_authenticated:
        return redirect('profile')  # Замените 'profile' на нужный URL

    if request.method == 'GET':
        # Рендерим страницу регистрации для GET-запросов
        return render(request, 'auth/register.html')

    if request.method == 'POST':
        # Получаем данные из формы
        email = request.POST.get('email')
        password = request.POST.get('password')
        password2 = request.POST.get('password2')
        surname = request.POST.get('familia')
        name = request.POST.get('name')
        patronymic = request.POST.get('otchestvo')
        tel = request.POST.get('tel')

        # Проверяем, совпадают ли пароли
        if password != password2:
            return JsonResponse({'success': False, 'message': 'Пароли не совпадают.'})

        # Проверяем, что все обязательные поля заполнены
        if not email or not password or not surname or not name or not tel:
            return JsonResponse({'success': False, 'message': 'Пожалуйста, заполните все обязательные поля.'})

        # Проверяем, что email уникален
        if User.objects.filter(email=email).exists():
            return JsonResponse({'success': False, 'message': 'Пользователь с такой почтой уже существует.'})

        # Регулярное выражение для проверки русских букв
        pattern = r'^[А-Яа-яЁё\s-]+$'

        # Проверяем, что ФИО содержит только кириллицу
        if not re.match(pattern, surname):
            return JsonResponse({'success': False, 'message': 'Фамилия должна содержать только русские буквы.'})

        if not re.match(pattern, name):
            return JsonResponse({'success': False, 'message': 'Имя должно содержать только русские буквы.'})

        if patronymic and not re.match(pattern, patronymic):  # Отчество может быть пустым
            return JsonResponse({'success': False, 'message': 'Отчество должно содержать только русские буквы.'})

        # Проверка сложности пароля
        if len(password) < 8:
            return JsonResponse({'success': False, 'message': 'Пароль должен содержать не менее 8 символов.'})

        if not re.search(r'[A-Z]', password):
            return JsonResponse({'success': False, 'message': 'Пароль должен содержать хотя бы одну заглавную букву.'})

        if not re.search(r'[a-z]', password):
            return JsonResponse({'success': False, 'message': 'Пароль должен содержать хотя бы одну строчную букву.'})

        if not re.search(r'[0-9]', password):
            return JsonResponse({'success': False, 'message': 'Пароль должен содержать хотя бы одну цифру.'})

        if not re.search(r'[!@#$%^&*(),.?":{}|<>]', password):
            return JsonResponse({'success': False, 'message': 'Пароль должен содержать хотя бы один специальный символ.'})

        # Создаем пользователя и запись в таблице Citizen
        try:
            user = User.objects.create_user(email=email, password=password)
            citizen = Citizen.objects.create(
                surname=surname,
                name=name,
                patronymic=patronymic,
                tel=tel,
                email=email,
                id_city=None,  # Город можно указать позже
                id_street=None,  # Улицу можно указать позже
                house="",  # Указываем пустую строку вместо NULL
                flat=0  # Указываем значение по умолчанию
            )
            user.id_citizen = citizen
            user.save()

            return JsonResponse({'success': True, 'message': 'Регистрация прошла успешно!'})

        except Exception as e:
            return JsonResponse({'success': False, 'message': f'Ошибка при регистрации: {str(e)}'})

    # Если метод запроса не GET или POST, возвращаем ошибку
    return JsonResponse({'success': False, 'message': 'Недопустимый метод запроса.'})



@login_required
def profile_view(request):
    return render(request, 'auth/profile.html')

def logout_view(request):
    logout(request)
    messages.success(request, 'Вы успешно вышли из системы.')
    return redirect('index')


# Городская служба

# Статистика
@login_required
def service_statistics(request):
    if not request.user.id_sotrudnik:
        messages.error(request, 'У вас нет доступа к этой странице.')
        return redirect('home')

    service = request.user.id_sotrudnik.id_service

    # Фильтры
    status_filter = request.GET.get('status')
    category_filter = request.GET.get('category')
    date_from = request.GET.get('date_from')
    date_to = request.GET.get('date_to')
    specific_date = request.GET.get('specific_date')

    # Базовый запрос для обращений службы
    appeals = Appeals.objects.filter(id_service=service)

    # Применение фильтров
    if category_filter:
        appeals = appeals.filter(id_category__id=category_filter)
    if date_from and date_to:
        appeals = appeals.filter(date_time__range=[date_from, date_to])
    if specific_date:
        appeals = appeals.filter(date_time__date=specific_date)

    # Создаем подзапрос для получения последних статусов
    latest_status_subquery = Processing_appeals.objects.filter(
        id_appeal=OuterRef('id')
    ).order_by('-date_time_setting_status').values('id_status')[:1]

    # Аннотируем обращения последним статусом
    appeals_with_last_status = appeals.annotate(
        last_status_id=Subquery(latest_status_subquery)
    ).annotate(
        last_status_name=Subquery(
            Status.objects.filter(id=OuterRef('last_status_id'))
            .values('name_status')[:1]
        )
    )

    # Фильтрация по последнему статусу (если указан фильтр)
    if status_filter:
        appeals_with_last_status = appeals_with_last_status.filter(
            last_status_id=status_filter)

    # Статистика по статусам
    status_stats = appeals_with_last_status.values('last_status_name').annotate(
        total=Count('id')).order_by('-total')

    # Статистика по категориям
    category_stats = appeals.values('id_category__name_official').annotate(
        total=Count('id')).order_by('-total')

    # Рассчитываем проценты
    total_appeals = appeals.count()
    for stat in status_stats:
        stat['percentage'] = (stat['total'] / total_appeals * 100) if total_appeals else 0
    for stat in category_stats:
        stat['percentage'] = (stat['total'] / total_appeals * 100) if total_appeals else 0

    context = {
        'status_stats': status_stats,
        'category_stats': category_stats,
        'statuses': Status.objects.all(),
        'categories': Category.objects.all(),
    }

    return render(request, 'service/service_statistics.html', context)


@login_required
def employee_appeals(request):
    if not request.user.id_sotrudnik:
        return redirect('profile')

    service = request.user.id_sotrudnik.id_service
    appeals = Appeals.objects.filter(id_service=service)

    # Фильтрация по статусу
    status_filter = request.GET.get('status')
    if status_filter:
        appeals = appeals.filter(processing_appeals__id_status=status_filter).distinct()

    # Фильтрация по дате (используем date_time вместо date_appeal)
    date_from = request.GET.get('date_from')
    date_to = request.GET.get('date_to')
    exact_date = request.GET.get('exact_date')

    if exact_date:
        appeals = appeals.filter(date_time__date=exact_date)
    else:
        if date_from:
            appeals = appeals.filter(date_time__date__gte=date_from)
        if date_to:
            appeals = appeals.filter(date_time__date__lte=date_to)

    # Получаем все возможные статусы для фильтра
    statuses = Status.objects.all()

    return render(request, 'service/appeals_service.html', {
        'appeals': appeals,
        'statuses': statuses,
        'current_status': status_filter,
        'date_from': date_from or '',
        'date_to': date_to or '',
        'exact_date': exact_date or '',
    })


@login_required
def view_appeal(request, appeal_id):
    """Просмотр и изменение статуса обращения"""
    appeal = get_object_or_404(Appeals, id=appeal_id)

    # Проверяем доступ сотрудника к обращению
    if appeal.id_service != request.user.id_sotrudnik.id_service:
        messages.error(request, 'У вас нет доступа к этому обращению.')
        return redirect('employee_appeals')

    # Получаем текущий статус обращения
    current_status = appeal.processing_appeals_set.last().id_status.name_status

    # Получаем список всех статусов, кроме "Принято"
    statuses = Status.objects.exclude(name_status='Принято')

    # Фильтруем доступные статусы в зависимости от текущего статуса
    if current_status == 'Принято':
        statuses = statuses.filter(name_status__in=['В работе', 'Отклонено'])
    elif current_status == 'В работе':
        statuses = statuses.filter(name_status__in=['Выполнено', 'Отклонено'])
    elif current_status in ['Выполнено', 'Отклонено']:
        statuses = statuses.none()

    completed_status_id = Status.objects.filter(name_status='Выполнено').values_list('id', flat=True).first()
    rejected_status_id = Status.objects.filter(name_status='Отклонено').values_list('id', flat=True).first()
    blocked_statuses = ['Выполнено', 'Отклонено']

    if request.method == 'POST':
        new_status_id = request.POST.get('status')
        if new_status_id:
            new_status = get_object_or_404(Status, id=new_status_id)

            # Проверка на статус "Выполнено"
            if new_status.name_status == 'Выполнено':
                # Проверяем наличие фото
                if not request.FILES.get('photo'):
                    messages.error(request, 'Для статуса "Выполнено" необходимо загрузить фото.')
                    return redirect('view_appeal', appeal_id=appeal.id)

                # Проверяем, что загруженный файл является изображением
                uploaded_file = request.FILES.get('photo')
                if uploaded_file:
                    allowed_types = ['image/jpeg', 'image/png', 'image/gif', 'image/bmp']
                    if uploaded_file.content_type not in allowed_types:
                        messages.error(request,
                                       'Пожалуйста, загрузите файл изображения (JPEG, PNG, GIF или BMP).')
                        return redirect('view_appeal', appeal_id=appeal.id)

            # Проверка на статус "Отклонено" и наличие причины
            if new_status.name_status == 'Отклонено' and not request.POST.get('rejection_reason'):
                messages.error(request, 'Для статуса "Отклонено" необходимо указать причину.')
                return redirect('view_appeal', appeal_id=appeal.id)

            # Создаем запись о статусе
            processing = Processing_appeals.objects.create(
                id_appeal=appeal,
                id_status=new_status,
                date_time_setting_status=timezone.now(),
                rejection_reason=request.POST.get('rejection_reason', ''),
            )

            # Если прикреплен файл, сохраняем его (только если это изображение)
            if 'photo' in request.FILES and new_status.name_status == 'Выполнено':
                processing.photo = request.FILES['photo']
                processing.save()

                # Проставляем сотрудника в обращении
                appeal.id_sotrudnik = request.user.id_sotrudnik
                appeal.save()

            messages.success(request, f'Статус обращения {appeal.id} изменён на {new_status.name_status}.')
            return redirect('view_appeal', appeal_id=appeal.id)

    status_history = Processing_appeals.objects.filter(id_appeal=appeal).order_by('-date_time_setting_status')

    return render(request, 'service/view_appeal.html', {
        'appeal': appeal,
        'statuses': statuses,
        'status_history': status_history,
        'completed_status_id': completed_status_id,
        'rejected_status_id': rejected_status_id,
        'blocked_statuses': blocked_statuses,
    })




@login_required
def create_report(request):
    if request.method == 'POST':
        form = ReportForm(request.POST)
        if form.is_valid():
            start_date = form.cleaned_data['start_date']
            end_date = form.cleaned_data['end_date']

            # Фильтруем обращения за выбранный период
            appeals = Appeals.objects.filter(date_time__range=(start_date, end_date))

            # Фильтруем выполненные обращения
            completed_appeals = Processing_appeals.objects.filter(
                id_status__name_status='Выполнено',
                date_time_setting_status__range=(start_date, end_date)
            )

            # Создаем Word-документ
            document = Document()
            document.add_heading('Отчет по обращениям', 0)

            # Добавляем информацию о периоде
            document.add_paragraph(f'Период: с {start_date} по {end_date}')

            # Добавляем статистику
            document.add_heading('Статистика', level=1)
            document.add_paragraph(f'Всего обращений: {appeals.count()}')
            document.add_paragraph(f'Выполнено обращений: {completed_appeals.count()}')

            # Добавляем таблицу с выполненными обращениями
            document.add_heading('Выполненные обращения', level=1)
            table = document.add_table(rows=1, cols=4)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Номер обращения'
            hdr_cells[1].text = 'Гражданин'
            hdr_cells[2].text = 'Описание'
            hdr_cells[3].text = 'Дата выполнения'

            for appeal in completed_appeals:
                row_cells = table.add_row().cells
                row_cells[0].text = str(appeal.id_appeal.id)
                row_cells[1].text = f"{appeal.id_appeal.id_sitizen.surname} {appeal.id_appeal.id_sitizen.name}"
                row_cells[2].text = appeal.id_appeal.description_problem
                row_cells[3].text = appeal.date_time_setting_status.strftime('%Y-%m-%d %H:%M:%S')

            # Сохраняем документ в байтовый поток
            buffer = io.BytesIO()
            document.save(buffer)
            buffer.seek(0)

            # Возвращаем документ как HTTP-ответ
            response = HttpResponse(
                buffer.getvalue(),
                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
            response['Content-Disposition'] = f'attachment; filename=report_{start_date}_to_{end_date}.docx'
            return response
    else:
        form = ReportForm()

    return render(request, 'service/create_report.html', {'form': form})


# _______________________


# Житель

# Функция для проверки ФИО (только кириллица)
def validate_cyrillic(text):
    return bool(re.fullmatch(r'^[А-ЯЁа-яё\s-]+$', text))


# Обновление данных пользователя
@login_required
def update_profile(request):
    user = request.user

    if request.method == 'POST':
        # Для жителей
        if user.id_citizen:
            return handle_citizen_update(request, user)
        # Для сотрудников
        elif user.id_sotrudnik:
            return handle_employee_update(request, user)
        else:
            messages.error(request, 'Неизвестный тип пользователя.')
            return redirect('update_profile')

    # Для GET запроса
    return render_update_profile_page(request, user)


def handle_citizen_update(request, user):
    citizen = user.id_citizen
    surname = request.POST.get('surname', '').strip()
    name = request.POST.get('name', '').strip()
    patronymic = request.POST.get('patronymic', '').strip()
    tel = request.POST.get('tel', '').strip()
    email = request.POST.get('email', '').strip()
    city_name = request.POST.get('city', '').strip()
    street_name = request.POST.get('street', '').strip()
    house = request.POST.get('house', '').strip()
    flat = request.POST.get('flat', '').strip()

    # Проверка ФИО (только кириллица)
    if not validate_cyrillic(surname) or not validate_cyrillic(name) or (
            patronymic and not validate_cyrillic(patronymic)):
        messages.error(request, 'Фамилия, имя и отчество должны содержать только кириллические буквы.')
        return redirect('update_profile')

    # Обновление данных
    citizen.surname = surname
    citizen.name = name
    citizen.patronymic = patronymic
    citizen.tel = tel
    citizen.email = email
    citizen.house = house
    citizen.flat = flat if flat else None

    if city_name:
        city, _ = City.objects.get_or_create(name_city=city_name)
        citizen.id_city = city
    if street_name:
        street, _ = Street.objects.get_or_create(name_street=street_name)
        citizen.id_street = street

    citizen.save()
    messages.success(request, 'Профиль успешно обновлён.')
    return redirect('update_profile')


def handle_employee_update(request, user):
    employee = user.id_sotrudnik
    surname = request.POST.get('surname', '').strip()
    name = request.POST.get('name', '').strip()
    patronymic = request.POST.get('patronymic', '').strip()

    # Проверка ФИО (только кириллица)
    if not validate_cyrillic(surname) or not validate_cyrillic(name) or (
            patronymic and not validate_cyrillic(patronymic)):
        messages.error(request, 'Фамилия, имя и отчество должны содержать только кириллические буквы.')
        return redirect('update_profile')

    # Обновление данных сотрудника
    employee.surname = surname
    employee.name = name
    employee.patronymic = patronymic
    employee.save()

    messages.success(request, 'Профиль успешно обновлён.')
    return redirect('update_profile')


def render_update_profile_page(request, user):
    context = {}

    if user.id_citizen:
        context['is_citizen'] = True
        context['profile_data'] = {
            'surname': user.id_citizen.surname,
            'name': user.id_citizen.name,
            'patronymic': user.id_citizen.patronymic,
            'tel': user.id_citizen.tel,
            'email': user.id_citizen.email,
            'city': user.id_citizen.id_city.name_city if user.id_citizen.id_city else '',
            'street': user.id_citizen.id_street.name_street if user.id_citizen.id_street else '',
            'house': user.id_citizen.house,
            'flat': user.id_citizen.flat,
        }
    elif user.id_sotrudnik:
        context['is_employee'] = True
        context['profile_data'] = {
            'surname': user.id_sotrudnik.surname,
            'name': user.id_sotrudnik.name,
            'patronymic': user.id_sotrudnik.patronymic,
        }

    return render(request, 'auth/update_profile.html', context)
# ______________________________

# Смена пароля
@login_required
def change_password(request):
    if request.method == 'POST':
        old_password = request.POST.get('old_password')
        new_password = request.POST.get('new_password')
        confirm_password = request.POST.get('confirm_password')

        # Проверка старого пароля
        if not request.user.check_password(old_password):
            return JsonResponse({'success': False, 'message': 'Старый пароль введен неверно.'})

        # Проверка совпадения новых паролей
        if new_password != confirm_password:
            return JsonResponse({'success': False, 'message': 'Новые пароли не совпадают.'})

        # Проверка сложности нового пароля
        if len(new_password) < 8:
            return JsonResponse({'success': False, 'message': 'Пароль должен содержать не менее 8 символов.'})

        if not re.search(r'[A-ZА-Я]', new_password):
            return JsonResponse({'success': False, 'message': 'Пароль должен содержать хотя бы одну заглавную букву.'})

        if not re.search(r'[a-zа-я]', new_password):
            return JsonResponse({'success': False, 'message': 'Пароль должен содержать хотя бы одну строчную букву.'})

        if not re.search(r'[0-9]', new_password):
            return JsonResponse({'success': False, 'message': 'Пароль должен содержать хотя бы одну цифру.'})

        if not re.search(r'[!@#$%^&*(),.?":{}|<>]', new_password):
            return JsonResponse({'success': False, 'message': 'Пароль должен содержать хотя бы один специальный символ.'})

        # Если все проверки пройдены - меняем пароль
        request.user.set_password(new_password)
        request.user.save()
        update_session_auth_hash(request, request.user)

        return JsonResponse({'success': True, 'message': 'Пароль успешно изменен!'})

    return render(request, 'auth/change_password.html')


# ОБращения
@login_required
def create_appeal(request):
    if request.method == 'POST':
        form = AppealForm(request.POST, request.FILES)
        if form.is_valid():
            # Создаем обращение, но не сохраняем его еще
            appeal = form.save(commit=False)
            appeal.id_sitizen = request.user.id_citizen
            appeal.date_time = timezone.now()

            # Сохраняем обращение, чтобы получить ID
            appeal.save()

            # Получаем статус "Принято" (ID = 1)
            status = Status.objects.get(id=1)  # Или name_status='Принято'

            # Создаем запись в Processing_appeals без фото
            Processing_appeals.objects.create(
                id_appeal=appeal,
                id_status=status,
                date_time_setting_status=timezone.now(),
            )

            # Загружаем фото только для модели Appeals
            if 'photo' in request.FILES:
                appeal.photo = request.FILES['photo']
                appeal.save()  # Сохраняем фото в объект Appeals

            return redirect('profile')
    else:
        form = AppealForm()
    return render(request, 'appeals/create_appeal.html', {'form': form})



@login_required
def view_appeals(request):
    # Получаем подзапрос с последним статусом для каждого обращения
    latest_status_subquery = Processing_appeals.objects.filter(
        id_appeal=OuterRef('id')
    ).order_by('-date_time_setting_status').values('id_status')[:1]

    # Джойн статусов и обращений
    appeals = Appeals.objects.filter(id_sitizen=request.user.id_citizen).annotate(
        latest_status_id=Subquery(latest_status_subquery)
    ).select_related('id_sitizen')

    # Присоединяем статусы к запросу через Python (чтобы убрать Subquery, если он не сработал)
    status_dict = {status.id: status.name_status for status in Status.objects.all()}
    for appeal in appeals:
        appeal.latest_status = status_dict.get(appeal.latest_status_id, "Нет статуса")

    # Фильтр по статусу
    status_id = request.GET.get('status')
    if status_id:
        appeals = [appeal for appeal in appeals if appeal.latest_status_id == int(status_id)]

    # Фильтр по дате
    date_filter = request.GET.get('date')
    if date_filter:
        try:
            date_obj = make_aware(datetime.strptime(date_filter, "%Y-%m-%d"))
            appeals = [appeal for appeal in appeals if appeal.date_time.date() == date_obj.date()]
        except ValueError:
            pass

    # Получаем все возможные статусы
    all_statuses = Status.objects.all()

    # Дебаг
    for appeal in appeals:
        print(f"Обращение {appeal.id}: последний статус — {appeal.latest_status}")

    return render(request, 'appeals/view_appeals.html', {
        'appeals': appeals,
        'all_statuses': all_statuses,
        'status_id': status_id,
        'date_filter': date_filter,
    })






@login_required
def appeal_detail(request, appeal_id):
    appeal = get_object_or_404(Appeals, id=appeal_id, id_sitizen=request.user.id_citizen)

    # Получаем историю статусов с фотографиями
    statuses = Processing_appeals.objects.filter(id_appeal=appeal).order_by('date_time_setting_status')

    # Получаем последний статус
    latest_status = statuses.last()
    latest_status_name = latest_status.id_status.name_status if latest_status else "Неизвестно"

    return render(request, 'appeals/appeal_detail.html', {
        'appeal': appeal,
        'statuses': statuses,
        'latest_status': latest_status_name
    })

@login_required
def edit_appeal(request, appeal_id):
    appeal = get_object_or_404(Appeals, id=appeal_id, id_sitizen=request.user.id_citizen)

    # Проверяем статус обращения
    latest_status = Processing_appeals.objects.filter(id_appeal=appeal).order_by('-date_time_setting_status').first()
    if latest_status and latest_status.id_status.name_status != "Принято":
        messages.error(request, "Обращение нельзя редактировать, так как оно уже в обработке.")
        return redirect('appeal_detail', appeal_id=appeal.id)

    if request.method == "POST":
        form = Edit_AppealForm(request.POST, request.FILES, instance=appeal)
        if form.is_valid():
            form.save()
            messages.success(request, "Обращение успешно обновлено.")
            return redirect('appeal_detail', appeal_id=appeal.id)
    else:
        form = Edit_AppealForm(instance=appeal)

    return render(request, 'appeals/edit_appeal.html', {'form': form, 'appeal': appeal})

@login_required
def delete_appeal(request, appeal_id):
    appeal = get_object_or_404(Appeals, id=appeal_id, id_sitizen=request.user.id_citizen)

    # Проверяем статус обращения
    latest_status = Processing_appeals.objects.filter(id_appeal=appeal).order_by('-date_time_setting_status').first()
    if latest_status and latest_status.id_status.name_status != "Принято":
        messages.error(request, "Обращение нельзя удалить, так как оно уже в обработке.")
        return redirect('appeal_detail', appeal_id=appeal.id)

    if request.method == "POST":
        appeal.delete()
        messages.success(request, "Обращение и его файлы удалены.")
        return redirect('view_appeals')

    return redirect('appeal_detail', appeal_id=appeal.id)



# ЧАТ
# В almet/views.py

# Пример передачи данных о сообщениях в шаблон
@login_required
def chat(request, appeal_id):
    appeal = get_object_or_404(Appeals, id=appeal_id)
    user = request.user

    # Определение имени пользователя
    if user.id_citizen:
        # Обычный пользователь (житель)
        sender_name = f"{user.id_citizen.surname} {user.id_citizen.name}"
    elif user.id_sotrudnik:
        # Сотрудник
        sender_name = f"{user.id_sotrudnik.id_service.name}"
    else:
        # Администратор
        sender_name = "Админ"

    print(f"User: {sender_name}")  # Для отладки
    # Получаем все сообщения для этого обращения
    chat_messages = Message.objects.filter(id_appeals=appeal).order_by('created_at')

    if request.method == "POST":
        message = request.POST.get("message")
        if message:
            send_message_to_chat(request, appeal_id, message, user)  # Отправка сообщения через WebSocket

    return render(request, 'chat/chat.html', {
        'appeal': appeal,
        'chat_messages': chat_messages,
        'user': user,
    })

@csrf_exempt
@login_required
def delete_message(request, message_id):
    try:
        message = Message.objects.get(id=message_id, sender=request.user)
        message.is_deleted = True
        message.save()

        # Получаем объект channel_layer
        channel_layer = get_channel_layer()

        # Отправляем сообщение в группу WebSocket
        async_to_sync(channel_layer.group_send)(
            f"chat_{message.id_appeals.id}",  # Имя группы
            {
                "type": "chat_message",
                "message": json.dumps({
                    "message_id": message.id,
                    "is_deleted": True,
                }),
            }
        )

        return JsonResponse({"success": True})
    except Message.DoesNotExist:
        return JsonResponse({"success": False, "error": "Сообщение не найдено"})
    except Exception as e:
        return JsonResponse({"success": False, "error": str(e)})

@csrf_exempt
@login_required
def edit_message(request, message_id):
    if request.method == 'POST':
        message = get_object_or_404(Message, id=message_id, sender=request.user)
        new_message = request.POST.get('message')
        message.message = new_message
        message.is_edited = True
        message.save()

        # Отправка обновления через WebSocket
        chat_socket_group = f'chat_{message.id_appeals.id}'
        message_data = {
            'message_id': message.id,
            'message': new_message,
            'is_edited': True
        }
        async_to_sync(get_channel_layer.group_send)(
            chat_socket_group,
            {
                'type': 'chat_message',
                'message': json.dumps(message_data)
            }
        )

        return JsonResponse({'success': True})
    return JsonResponse({'success': False})

# Для WebSocket отправки сообщений
def send_message_to_chat(request, appeal_id, message, sender):
    sender = request.user
    # Формируем имя отправителя
    if sender.id_citizen:
        sender_name = f"{sender.id_citizen.surname} {sender.id_citizen.name}"
    elif sender.id_sotrudnik:
        sender_name = f"{sender.id_sotrudnik.id_service.name}"
    else:
        sender_name = "Админ"

    # Создание сообщения в базе данных
    chat_message = Message.objects.create(
        id_appeals=Appeals.objects.get(id=appeal_id),
        sender=sender,
        message=message,
    )

    # Если есть изображение, сохраняем его
    if 'image' in request.FILES:
        image_file = request.FILES['image']
        chat_message.image.save(image_file.name, image_file, save=True)

    # Отправка сообщения всем подключенным пользователям через WebSocket
    chat_socket_group = f'chat_{appeal_id}'
    message_data = {
        'sender_id': sender.id,
        'sender_name': sender_name,  # Передаем корректное имя
        'message': message,
        'image_url': chat_message.image.url if chat_message.image else None
    }
    async_to_sync(get_channel_layer.group_send)(
        chat_socket_group,
        {
            'type': 'chat_message',
            'message': json.dumps(message_data)
        }
    )



# ______________________________________________

# Администратор
# Проверка, является ли пользователь администратором
def is_admin(user):
    return user.is_staff or user.is_superuser

# Страница управления категориями
@login_required
@user_passes_test(is_admin)
def admin_categories(request):
    search_query = request.GET.get('search', '')  # Получаем поисковый запрос из GET-параметра

    # Если запрос AJAX, отправляем JSON-ответ с результатами поиска
    if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
        categories = Category.objects.filter(
            Q(name_official__icontains=search_query) | Q(name_short__icontains=search_query)
        ) if search_query else Category.objects.all()

        data = [
            {
                'id': category.id,
                'name_official': category.name_official,
                'name_short': category.name_short
            }
            for category in categories
        ]
        return JsonResponse({'categories': data})

    # Если обычный запрос, отдаем страницу с категориями
    categories = Category.objects.all()

    if request.method == 'POST':
        # Добавление категории
        if 'add_category' in request.POST:
            name_official = request.POST.get('name_official')
            name_short = request.POST.get('name_short')
            if name_official:
                Category.objects.create(name_official=name_official, name_short=name_short)
                messages.success(request, 'Категория успешно добавлена.')
            else:
                messages.error(request, 'Официальное название обязательно.')

        # Загрузка файла (XML или Excel)
        elif 'upload_file' in request.POST:
            uploaded_file = request.FILES.get('file')
            if uploaded_file:
                file_name = uploaded_file.name.lower()  # Приводим имя файла к нижнему регистру
                try:
                    if file_name.endswith('.xml'):
                        # Обработка XML-файла
                        tree = ET.parse(uploaded_file)
                        root = tree.getroot()

                        # Проверяем, что корневой элемент - <categories>
                        if root.tag != 'categories':
                            messages.error(request, 'Некорректный формат XML: ожидается корневой элемент <categories>.')
                            return redirect('admin_categories')

                        # Обрабатываем каждую категорию
                        for item in root.findall('category'):
                            name_official = item.find('name_official')
                            name_short = item.find('name_short')

                            # Проверяем, что элемент <name_official> существует и не пустой
                            if name_official is not None and name_official.text:
                                # Создаем категорию, если название не пустое
                                Category.objects.create(
                                    name_official=name_official.text,
                                    name_short=name_short.text if name_short is not None else None
                                )
                            else:
                                messages.warning(request, 'Найдена категория без названия. Пропущено.')

                        messages.success(request, 'Категории успешно загружены из XML.')

                    elif file_name.endswith(('.xlsx', '.xls')):  # Обработка Excel-файла
                        df = pd.read_excel(uploaded_file)

                        if 'name_official' not in df.columns:
                            messages.error(request, 'В файле отсутствует колонка "name_official".')
                            return redirect('admin_categories')

                        for index, row in df.iterrows():
                            name_official = row['name_official']
                            name_short = row.get('name_short', None)

                            if pd.notna(name_official):
                                Category.objects.create(
                                    name_official=name_official,
                                    name_short=name_short
                                )
                            else:
                                messages.warning(request, f'Найдена категория без названия в строке {index + 1}. Пропущено.')

                        messages.success(request, 'Категории успешно загружены из Excel.')

                    else:
                        messages.error(request, 'Неподдерживаемый формат файла. Разрешены только XML, XLSX или XLS.')

                except ET.ParseError as e:
                    messages.error(request, f'Ошибка при разборе XML: {str(e)}')
                except pd.errors.EmptyDataError:
                    messages.error(request, 'Файл Excel пуст или имеет неправильный формат.')
                except Exception as e:
                    messages.error(request, f'Ошибка при обработке файла: {str(e)}')
            else:
                messages.error(request, 'Файл не выбран.')

        # Массовое удаление
        elif 'mass_delete' in request.POST:
            selected_categories = request.POST.getlist('selected_categories')
            if selected_categories:
                Category.objects.filter(id__in=selected_categories).delete()
                messages.success(request, 'Выбранные категории успешно удалены.')
            else:
                messages.error(request, 'Не выбрано ни одной категории для удаления.')

    return render(request, 'admin/categories.html', {'categories': categories, 'search_query': search_query})



# Администратор:: Изменение данных категорий
@login_required
@user_passes_test(is_admin)
def edit_category(request, category_id):
    category = get_object_or_404(Category, id=category_id)

    if request.method == 'POST':
        # Получаем данные из формы
        name_official = request.POST.get('name_official')
        name_short = request.POST.get('name_short')

        # Обновляем данные категории
        if name_official:
            category.name_official = name_official
            category.name_short = name_short
            category.save()
            messages.success(request, 'Категория успешно изменена.')
            return redirect('admin_categories')
        else:
            messages.error(request, 'Официальное название обязательно.')

    return render(request, 'admin/edit_category.html', {'category': category})


@login_required
@user_passes_test(is_admin)
def delete_category(request, category_id):
    category = get_object_or_404(Category, id=category_id)
    if request.method == 'POST':
        category.delete()
        messages.success(request, 'Категория успешно удалена.')
    return redirect('admin_categories')




# Страница управления пользователями
@login_required
@user_passes_test(is_admin)
def admin_users(request):
    users = User.objects.all()
    services = Service.objects.all()

    # Фильтр по роли
    role_filter = request.GET.get('role')
    if role_filter == 'citizen':
        users = users.filter(id_sotrudnik__isnull=True, is_staff=False, is_superuser=False)
    elif role_filter == 'employee':
        users = users.filter(id_sotrudnik__isnull=False)

    # Фильтр по службе
    service_filter = request.GET.get('service')
    if service_filter:
        users = users.filter(id_sotrudnik__id_service=service_filter)

    # Поиск по ФИО
    search_query = request.GET.get('search')
    if search_query:
        users = users.filter(
            Q(id_citizen__surname__icontains=search_query) |
            Q(id_citizen__name__icontains=search_query) |
            Q(id_citizen__patronymic__icontains=search_query) |
            Q(id_sotrudnik__surname__icontains=search_query) |
            Q(id_sotrudnik__name__icontains=search_query) |
            Q(id_sotrudnik__patronymic__icontains=search_query)
        )

    return render(request, 'admin/users.html', {
        'users': users,
        'services': services,
        'role_filter': role_filter,
        'service_filter': service_filter,
        'search_query': search_query,
    })


@login_required
@user_passes_test(is_admin)
def change_role(request, user_id):
    user = get_object_or_404(User, id=user_id)
    services = Service.objects.all()

    if request.method == 'POST':
        role = request.POST.get('role')
        service_id = request.POST.get('service')

        if role == 'employee':
            # Превращаем пользователя в сотрудника
            if not user.id_sotrudnik:
                service = get_object_or_404(Service, id=service_id)
                employee = Sotrudniki.objects.create(
                    surname=user.id_citizen.surname if user.id_citizen else 'Unknown',
                    name=user.id_citizen.name if user.id_citizen else 'Unknown',
                    id_service=service
                )
                user.id_sotrudnik = employee
                user.save()
                messages.success(request, f'Пользователь {user.email} теперь сотрудник.')
        elif role == 'user':
            # Превращаем сотрудника в пользователя
            if user.id_sotrudnik:
                user.id_sotrudnik.delete()
                user.id_sotrudnik = None
                user.save()
                messages.success(request, f'Сотрудник {user.email} теперь пользователь.')

        return redirect('admin_users')

    return render(request, 'admin/change_role.html', {
        'user': user,
        'services': services,
    })

@login_required
@user_passes_test(is_admin)
def delete_user(request, user_id):
    if request.method == 'POST':
        user = get_object_or_404(User, id=user_id)

        # Удаляем связанного жителя, если он существует
        if user.id_citizen:
            user.id_citizen.delete()

        # Удаляем связанного сотрудника, если он существует
        if user.id_sotrudnik:
            user.id_sotrudnik.delete()

        # Удаляем самого пользователя
        user.delete()

        messages.success(request, 'Пользователь и связанные данные успешно удалены.')
        return redirect('admin_users')  # Перенаправляем на страницу управления пользователями

    messages.error(request, 'Неверный запрос.')
    return redirect('admin_users')


# Страница создания городской службы на сайте
@login_required
@user_passes_test(is_admin)
def admin_create_service(request):
    if request.method == 'POST':
        form = ServiceForm(request.POST)
        if form.is_valid():
            service = form.save(commit=False)
            street_name = form.cleaned_data['street']  # Получаем введенное название улицы
            street, _ = Street.objects.get_or_create(name_street=street_name)  # Ищем или создаем улицу
            service.id_street = street  # Присваиваем объект Street, а не строку
            service.save()
            messages.success(request, 'Городская служба успешно создана.')
            return redirect('admin_create_service')
    else:
        form = ServiceForm()

    services = Service.objects.all()
    return render(request, 'admin/create_service.html', {'form': form, 'services': services})



@login_required
@user_passes_test(is_admin)
def admin_service_list(request):
    services = Service.objects.all()  # Получаем все городские службы
    return render(request, 'admin/create_service.html', {'services': services})


@login_required
@user_passes_test(is_admin)
def admin_edit_service(request, service_id):
    service = get_object_or_404(Service, id=service_id)  # Получаем службу по ID

    if request.method == 'POST':
        form = ServiceForm(request.POST, instance=service)  # Используем instance для редактирования
        if form.is_valid():
            form.save()
            messages.success(request, 'Городская служба успешно обновлена.')
            return redirect('admin_create_service')  # Перенаправляем на список служб
    else:
        form = ServiceForm(instance=service)  # Передаём существующую службу в форму

    return render(request, 'admin/edit_service.html', {'form': form, 'service': service})


@login_required
@user_passes_test(is_admin)
def admin_delete_service(request, service_id):
    service = get_object_or_404(Service, id=service_id)
    service.delete()
    messages.success(request, 'Городская служба успешно удалена.')
    return redirect('admin_create_service')



# Страница создания сотрудника городской службы
@login_required
@user_passes_test(is_admin)
def admin_create_employee(request):
    if request.method == 'POST':
        form = EmployeeRegistrationForm(request.POST)
        if form.is_valid():
            email = form.cleaned_data['email']
            surname = form.cleaned_data['surname']
            name = form.cleaned_data['name']
            patronymic = form.cleaned_data['patronymic']
            service = form.cleaned_data['service']

            # Генерируем случайный пароль
            alphabet = string.ascii_letters + string.digits
            password = ''.join(secrets.choice(alphabet) for i in range(12))  # Пароль из 12 символов


            # Создаем пользователя и запись в таблице Sotrudniki
            try:
                # Создаем пользователя
                user = User.objects.create_user(email=email, password=password)  # Передаем password, а не hashed_password

                # Создаем запись в таблице Sotrudniki
                employee = Sotrudniki.objects.create(
                    surname=surname,
                    name=name,
                    patronymic=patronymic,
                    id_service=service
                )

                # Связываем пользователя с записью в таблице Sotrudniki
                user.id_sotrudnik = employee
                user.save()

                # Создаем текстовый файл с паролем
                file_content = f"Email: {email}\nФамилия: {surname}\nИмя: {name}\nПароль: {password}"
                response = HttpResponse(file_content, content_type='text/plain')
                response['Content-Disposition'] = f'attachment; filename="employee_{employee.id}_password.txt"'

                # Авторизуем пользователя
                messages.success(request, 'Регистрация сотрудника прошла успешно! Пароль скачан.')
                return response
            except Exception as e:
                messages.error(request, f'Ошибка при регистрации: {str(e)}')
                return render(request, 'auth/register_employee.html', {'form': form})
    else:
        form = EmployeeRegistrationForm()
    return render(request, 'admin/create_employee.html', {'form': form})



# Страница просмотра всех обращений
@login_required
@user_passes_test(is_admin)
def admin_all_appeals(request):
    appeals = Appeals.objects.all()
    services = Service.objects.all()
    statuses = Status.objects.all()

    # Подзапрос для получения последнего статуса
    latest_status_subquery = Processing_appeals.objects.filter(
        id_appeal=OuterRef('pk')
    ).order_by('-date_time_setting_status').values('id_status')[:1]

    # Аннотируем обращения последним статусом
    appeals = appeals.annotate(
        latest_status=Subquery(latest_status_subquery)
    )

    # Фильтрация по дате
    start_date = request.GET.get('start_date')
    end_date = request.GET.get('end_date')
    if start_date and end_date:
        appeals = appeals.filter(date_time__range=[start_date, end_date])

    # Фильтрация по статусу
    status_id = request.GET.get('status')
    if status_id:
        appeals = appeals.filter(latest_status=status_id)

    # Фильтрация по службе
    service_id = request.GET.get('service')
    if service_id:
        appeals = appeals.filter(id_service=service_id)

    # Убираем дубликаты (на всякий случай)
    appeals = appeals.distinct()

    return render(request, 'admin/all_appeals.html', {
        'appeals': appeals,
        'services': services,
        'statuses': statuses,
    })

@login_required
@user_passes_test(is_admin)
def assign_service(request, appeal_id):
    appeal = get_object_or_404(Appeals, id=appeal_id)
    if request.method == 'POST':
        service_id = request.POST.get('service')
        if service_id:
            service = get_object_or_404(Service, id=service_id)
            appeal.id_service = service
            appeal.save()
            messages.success(request, f'Обращение {appeal.id} назначено на службу {service.name}.')
        else:
            appeal.id_service = None
            appeal.save()
            messages.success(request, f'Обращение {appeal.id} больше не назначено на службу.')
    return redirect('admin_all_appeals')


@login_required
@user_passes_test(is_admin)
def admin_view_appeal(request, appeal_id):
    appeal = get_object_or_404(Appeals, id=appeal_id)
    statuses = Status.objects.all()
    services = Service.objects.all()

    if request.method == 'POST':
        if 'status' in request.POST:
            status_id = request.POST.get('status')
            status = get_object_or_404(Status, id=status_id)
            Processing_appeals.objects.create(
                id_appeal=appeal,
                id_status=status,
                date_time_setting_status=now()  # Добавляем текущее время
            )
            messages.success(request, f'Статус обращения {appeal.id} обновлён.')

        elif 'service' in request.POST:
            service_id = request.POST.get('service')
            if service_id:
                service = get_object_or_404(Service, id=service_id)
                appeal.id_service = service
            else:
                appeal.id_service = None
            appeal.save()
            messages.success(request, f'Обращение {appeal.id} обновлено.')

        elif 'delete' in request.POST:
            appeal.delete()
            messages.success(request, 'Обращение удалено.')
            return redirect('admin_all_appeals')

    return render(request, 'admin/view_appeal.html', {
        'appeal': appeal,
        'statuses': statuses,
        'services': services,
    })

# Статистика у админа
@login_required
def admin_statistics(request):
    if not request.user.is_staff:
        messages.error(request, 'У вас нет доступа к этой странице.')
        return redirect('home')

    # Фильтры
    status_filter = request.GET.get('status')
    category_filter = request.GET.get('category')
    service_filter = request.GET.get('service')
    date_from = request.GET.get('date_from')
    date_to = request.GET.get('date_to')
    specific_date = request.GET.get('specific_date')

    # Базовый запрос для обращений
    appeals = Appeals.objects.all()

    # Применение фильтров
    if category_filter:
        appeals = appeals.filter(id_category__id=category_filter)
    if service_filter:
        appeals = appeals.filter(id_service__id=service_filter)
    if date_from and date_to:
        appeals = appeals.filter(date_time__range=[date_from, date_to])
    if specific_date:
        appeals = appeals.filter(date_time__date=specific_date)

    # Создаем подзапрос для получения последних статусов
    latest_status_subquery = Processing_appeals.objects.filter(
        id_appeal=OuterRef('id')
    ).order_by('-date_time_setting_status').values('id_status')[:1]

    # Аннотируем обращения последним статусом
    appeals_with_last_status = appeals.annotate(
        last_status_id=Subquery(latest_status_subquery)
    ).annotate(
        last_status_name=Subquery(
            Status.objects.filter(id=OuterRef('last_status_id'))
            .values('name_status')[:1]
        )
    )

    # Фильтрация по последнему статусу (если указан фильтр)
    if status_filter:
        appeals_with_last_status = appeals_with_last_status.filter(
            last_status_id=status_filter)

    # Статистика по службам
    service_stats = appeals_with_last_status.values('id_service__name').annotate(
        total=Count('id')).order_by('-total')

    # Статистика по статусам
    status_stats = appeals_with_last_status.values('last_status_name').annotate(
        total=Count('id')).order_by('-total')

    # Статистика по категориям
    category_stats = appeals.values('id_category__name_official').annotate(
        total=Count('id')).order_by('-total')

    # Рассчитываем проценты
    total_appeals = appeals.count()
    for stat in service_stats:
        stat['percentage'] = (stat['total'] / total_appeals * 100) if total_appeals else 0
    for stat in status_stats:
        stat['percentage'] = (stat['total'] / total_appeals * 100) if total_appeals else 0
    for stat in category_stats:
        stat['percentage'] = (stat['total'] / total_appeals * 100) if total_appeals else 0

    context = {
        'service_stats': service_stats,
        'status_stats': status_stats,
        'category_stats': category_stats,
        'statuses': Status.objects.all(),
        'categories': Category.objects.all(),
        'services': Service.objects.all(),
    }

    return render(request, 'admin/admin_statistics.html', context)


